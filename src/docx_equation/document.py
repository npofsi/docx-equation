"""Developer-friendly helpers for writing equation DOCX files with python-docx."""

from __future__ import annotations

from pathlib import Path
import tempfile
from typing import Iterator

from docx import Document

from docx_equation.api import embed_mathml_placeholders
from docx_equation.shared.models import ConversionSummary, EquationSpec, OptionsLike


class EquationRegistry:
    """Collect MathML equations while composing a ``python-docx`` document.

    The registry writes hidden temporary marker runs into the in-memory
    document and replaces them during ``save``. Calling code does not need to
    create or track placeholder strings directly.
    """

    def __init__(self, prefix: str = "DOCX_EQ") -> None:
        self.prefix = prefix
        self._equations: list[EquationSpec] = []

    @property
    def equations(self) -> list[EquationSpec]:
        """Return a copy of the collected equation specifications."""
        return list(self._equations)

    def __len__(self) -> int:
        return len(self._equations)

    def __iter__(self) -> Iterator[EquationSpec]:
        return iter(self._equations)

    def add_inline(self, paragraph, mathml: str) -> EquationSpec:
        """Add an inline equation marker to an existing paragraph."""
        return self.add(paragraph, mathml, display=False)

    def add_display(self, document_or_paragraph, mathml: str, number: int | str | None = None, style: str | None = None) -> EquationSpec:
        """Add a display equation marker to a document or existing paragraph."""
        paragraph = (
            document_or_paragraph
            if hasattr(document_or_paragraph, "add_run")
            else document_or_paragraph.add_paragraph(style=style)
        )
        return self.add(paragraph, mathml, display=True, number=number)

    def add(
        self,
        paragraph,
        mathml: str,
        *,
        display: bool = False,
        number: int | str | None = None,
    ) -> EquationSpec:
        """Add an equation marker run and return its collected specification."""
        placeholder = self._next_placeholder()
        run = paragraph.add_run(placeholder)
        run.font.hidden = True
        spec = EquationSpec(placeholder=placeholder, mathml=mathml, display=display, number=number)
        self._equations.append(spec)
        return spec

    def save(
        self,
        document,
        output_path: str | Path,
        options: OptionsLike = None,
        work_dir: str | Path | None = None,
    ) -> ConversionSummary:
        """Save ``document`` and embed all registered equations."""
        target = Path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        if not self._equations:
            document.save(target)
            return ConversionSummary()

        with tempfile.TemporaryDirectory(prefix="docx_equation_registry_") as tmp:
            base_docx = Path(tmp) / "base.docx"
            document.save(base_docx)
            embed_work_dir = Path(work_dir) if work_dir is not None else Path(tmp) / "work"
            return embed_mathml_placeholders(base_docx, target, self.equations, options, embed_work_dir)

    def _next_placeholder(self) -> str:
        return f"{{{{{self.prefix}_{len(self._equations) + 1:03d}}}}}"


class EquationDocument:
    """Small wrapper combining a ``python-docx`` document and an equation registry."""

    def __init__(self, document=None, prefix: str = "DOCX_EQ") -> None:
        self.document = document if document is not None else Document()
        self.equations = EquationRegistry(prefix=prefix)

    def add_inline_equation(self, paragraph, mathml: str) -> EquationSpec:
        """Add an inline equation to ``paragraph``."""
        return self.equations.add_inline(paragraph, mathml)

    def add_display_equation(self, mathml: str, number: int | str | None = None, style: str | None = None) -> EquationSpec:
        """Add a display equation paragraph."""
        return self.equations.add_display(self.document, mathml, number=number, style=style)

    def save(self, output_path: str | Path, options: OptionsLike = None, work_dir: str | Path | None = None) -> ConversionSummary:
        """Save the document and embed registered equations."""
        return self.equations.save(self.document, output_path, options, work_dir)

    def __getattr__(self, name: str):
        return getattr(self.document, name)
