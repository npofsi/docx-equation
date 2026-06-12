"""DOCX embedding support for MathType-compatible equation objects."""

from __future__ import annotations

from pathlib import Path
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree
from PIL import Image, ImageDraw

from docx_equation.mathtype.ooxml import make_object_run
from docx_equation.omml import mathml_to_omml
from docx_equation.shared.mathml import parse_mathml, render_mathml_files
from docx_equation.mathtype.mtef import encode_mtef
from docx_equation.mathtype.ole import build_mathtype_ole_object
from docx_equation.shared.models import (
    ConversionSummary,
    EquationSpec,
    ExportOptions,
    MathTypeOptions,
    OptionsLike,
    normalize_options,
)
from docx_equation.shared.numbering import make_tabbed_equation_paragraph
from docx_equation.shared.ooxml import (
    IMAGE_REL,
    NS,
    OLE_REL,
    ancestor_paragraph,
    add_relationship,
    document_text_width_dxa,
    ensure_default,
    find_placeholder_run,
)


REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
DXEQ_NS = "https://github.com/npofsi/docx-equation/mathtype"
W_NS = NS["w"]


def embed_mathml_placeholders(
    input_docx: str | Path,
    output_docx: str | Path,
    equations: list[EquationSpec],
    options: OptionsLike = None,
    work_dir: str | Path | None = None,
) -> ConversionSummary:
    """Replace DOCX placeholders with MathType OLE objects generated from MathML."""
    opts = normalize_options(options, target="mathtype")
    source = Path(input_docx)
    target = Path(output_docx)
    target.parent.mkdir(parents=True, exist_ok=True)
    if work_dir is None:
        with tempfile.TemporaryDirectory(prefix="docx_equation_embed_") as tmp:
            return _embed(source, target, equations, opts, Path(tmp))
    return _embed(source, target, equations, opts, Path(work_dir))


def build_equation_docx(
    equations: list[str],
    output_docx: str | Path,
    options: OptionsLike = None,
) -> ConversionSummary:
    """Build a new DOCX containing MathType display equations."""
    opts = normalize_options(options, target="mathtype")
    target = Path(output_docx)
    target.parent.mkdir(parents=True, exist_ok=True)
    specs: list[EquationSpec] = []
    with tempfile.TemporaryDirectory(prefix="docx_equation_docx_") as tmp:
        base = Path(tmp) / "base.docx"
        doc = Document()
        for index, mathml in enumerate(equations, 1):
            placeholder = f"{{{{DOCX_EQ_{index:03d}}}}}"
            paragraph = doc.add_paragraph()
            paragraph.add_run(placeholder)
            specs.append(EquationSpec(placeholder=placeholder, mathml=mathml, display=True, number=index))
        doc.save(base)
        return embed_mathml_placeholders(base, target, specs, opts, Path(tmp) / "work")


def _embed(
    source: Path,
    target: Path,
    equations: list[EquationSpec],
    opts: ExportOptions,
    work_dir: Path,
) -> ConversionSummary:
    work_dir.mkdir(parents=True, exist_ok=True)
    mathml_dir = work_dir / "mathml"
    preview_dir = work_dir / "preview_png"
    ole_dir = work_dir / "ole"
    for directory in (mathml_dir, preview_dir, ole_dir):
        directory.mkdir(parents=True, exist_ok=True)

    summary = ConversionSummary(found=len(equations))
    parser = etree.XMLParser(resolve_entities=False, recover=True, remove_blank_text=False)
    with ZipFile(source) as zin:
        document_root = etree.fromstring(zin.read("word/document.xml"), parser)
        rels_root = etree.fromstring(zin.read("word/_rels/document.xml.rels"), parser)
        content_types_root = etree.fromstring(zin.read("[Content_Types].xml"), parser)

        document_root = _ensure_mc(document_root)
        ensure_default(content_types_root, "bin", "application/vnd.openxmlformats-officedocument.oleObject")
        ensure_default(content_types_root, "png", "image/png")

        for index, equation in enumerate(equations, 1):
            (mathml_dir / f"equation_{index:03d}.mml").write_text(equation.mathml, encoding="utf-8")

        mt_opts = opts.mathtype
        _render_previews(mathml_dir, preview_dir, mt_opts)

        media_entries: list[tuple[str, bytes]] = []
        ole_entries: list[tuple[str, bytes]] = []

        for index, equation in enumerate(equations, 1):
            try:
                placeholder_run = find_placeholder_run(document_root, equation.placeholder)
                preview_path = preview_dir / f"equation_{index:03d}.png"
                width_px, height_px = Image.open(preview_path).size

                image_name = f"mathtype_preview_{index:03d}.png"
                ole_name = f"oleObjectMathType{index:03d}.bin"
                image_rel_id = add_relationship(rels_root, IMAGE_REL, f"media/{image_name}")
                ole_rel_id = add_relationship(rels_root, OLE_REL, f"embeddings/{ole_name}")

                mtef = encode_mtef(parse_mathml(equation.mathml), mt_opts.mathtype_version)
                ole_bytes = build_mathtype_ole_object(mtef, mt_opts.prog_id)
                ole_entries.append((f"word/embeddings/{ole_name}", ole_bytes))
                media_entries.append((f"word/media/{image_name}", preview_path.read_bytes()))
                (ole_dir / ole_name).write_bytes(ole_bytes)

                replacement = make_object_run(
                    image_rel_id,
                    ole_rel_id,
                    width_px,
                    height_px,
                    index=index,
                    display=equation.display,
                    inline_height_pt=mt_opts.inline_height_pt,
                    display_height_pt=mt_opts.display_height_pt,
                    max_width_pt=mt_opts.max_width_pt,
                    preview_pt_per_px=mt_opts.preview_pt_per_px,
                    vertical_align="middle" if equation.display else None,
                    prog_id=mt_opts.prog_id,
                )
                if mt_opts.embed_mode == "alternate-content":
                    fallback = mathml_to_omml(
                        equation.mathml,
                        display=equation.display and equation.number is None,
                        options=opts.omml,
                    )
                    replacement = _alternate_content(replacement, fallback)
                if equation.display and equation.number is not None and opts.display_layout == "tabbed":
                    paragraph = make_tabbed_equation_paragraph(
                        replacement,
                        equation.number,
                        document_text_width_dxa(document_root),
                        numbering=opts.numbering,
                        style=opts.style,
                    )
                    placeholder_paragraph = ancestor_paragraph(placeholder_run)
                    if placeholder_paragraph is None:
                        raise ValueError(f"Placeholder paragraph was not found: {equation.placeholder}")
                    paragraph.tail = placeholder_paragraph.tail
                    placeholder_paragraph.getparent().replace(placeholder_paragraph, paragraph)
                else:
                    replacement.tail = placeholder_run.tail
                    placeholder_run.getparent().replace(placeholder_run, replacement)
                summary.converted += 1
            except Exception as exc:  # noqa: BLE001 - conversion continues per equation.
                summary.add_error(index, equation.placeholder, str(exc))

        document_xml = etree.tostring(document_root, encoding="utf-8", xml_declaration=True, standalone=True)
        rels_xml = etree.tostring(rels_root, encoding="utf-8", xml_declaration=True, standalone=True)
        content_types_xml = etree.tostring(content_types_root, encoding="utf-8", xml_declaration=True, standalone=True)

        with ZipFile(target, "w", ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                if info.filename == "word/document.xml":
                    zout.writestr(info, document_xml)
                elif info.filename == "word/_rels/document.xml.rels":
                    zout.writestr(info, rels_xml)
                elif info.filename == "[Content_Types].xml":
                    zout.writestr(info, content_types_xml)
                else:
                    zout.writestr(info, zin.read(info.filename))
            for name, data in media_entries + ole_entries:
                zout.writestr(name, data)
    return summary


def _render_previews(mathml_dir: Path, preview_dir: Path, opts: MathTypeOptions) -> None:
    try:
        render_mathml_files(mathml_dir, preview_dir, chrome_path=opts.chrome_path, font_px=opts.preview_font_px)
    except Exception:
        if opts.embed_mode == "png-preview":
            raise
        for mathml_path in sorted(mathml_dir.glob("equation_*.mml")):
            _write_fallback_preview(preview_dir / f"{mathml_path.stem}.png")


def _write_fallback_preview(path: Path) -> None:
    image = Image.new("RGB", (420, 96), "white")
    draw = ImageDraw.Draw(image)
    draw.text((16, 34), "MathType equation", fill="black")
    image.save(path)


def _alternate_content(choice_child: etree._Element, fallback_child: etree._Element) -> etree._Element:
    alternate = etree.Element(f"{{{MC_NS}}}AlternateContent", nsmap={"mc": MC_NS, "dxeq": DXEQ_NS})
    choice = etree.SubElement(alternate, f"{{{MC_NS}}}Choice", {"Requires": "dxeq"})
    choice.append(choice_child)
    fallback = etree.SubElement(alternate, f"{{{MC_NS}}}Fallback")
    fallback.append(fallback_child)
    return alternate


def _ensure_mc(root: etree._Element) -> etree._Element:
    if root.nsmap.get("dxeq") != DXEQ_NS:
        nsmap = dict(root.nsmap)
        nsmap["dxeq"] = DXEQ_NS
        replacement = etree.Element(root.tag, nsmap=nsmap)
        replacement.text = root.text
        replacement.tail = root.tail
        for name, value in root.attrib.items():
            replacement.set(name, value)
        replacement[:] = list(root)
        root = replacement

    ignorable_attr = f"{{{MC_NS}}}Ignorable"
    current = root.get(ignorable_attr, "")
    values = {item for item in current.split() if item}
    values.add("dxeq")
    root.set(ignorable_attr, " ".join(sorted(values)))
    return root
