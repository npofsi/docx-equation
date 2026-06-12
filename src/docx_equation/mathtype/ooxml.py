from __future__ import annotations

from pathlib import Path
import tempfile
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from lxml import etree

from docx_equation.mathtype.mtef import encode_mtef
from docx_equation.mathtype.ole import build_mathtype_ole_object
from docx_equation.mathtype.preview import make_preview_png
from docx_equation.shared.latex import parse_latex_subset
from docx_equation.shared.models import EquationStyle, NumberingOptions
from docx_equation.shared.numbering import make_tabbed_equation_paragraph


NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "o": "urn:schemas-microsoft-com:office:office",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
}

REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
IMAGE_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
OLE_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject"
PLACEHOLDER = "{{DOCX_EQUATION_OLE_1}}"


def build_demo_docx(formula: str, output_docx: str | Path, mathtype_version: str = "DSMT4") -> Path:
    target = Path(output_docx)
    target.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="docx_equation_docx_") as tmp:
        tmp_dir = Path(tmp)
        base_docx = tmp_dir / "base.docx"
        preview_path = tmp_dir / "preview.png"
        ole_path = tmp_dir / "oleObject1.bin"

        expr = parse_latex_subset(formula)
        ole_path.write_bytes(build_mathtype_ole_object(encode_mtef(expr, mathtype_version), f"Equation.{mathtype_version}"))
        preview_width, preview_height = make_preview_png(formula, preview_path)

        doc = Document()
        doc.add_heading("DOCX Equation Demo", level=1)
        paragraph = doc.add_paragraph()
        paragraph.add_run("Generated Equation.DSMT4 object: ")
        paragraph.add_run(PLACEHOLDER)
        doc.add_paragraph("This DOCX embeds a MathType-compatible OLE object and uses a PNG preview for display.")
        doc.save(base_docx)

        _inject_object(base_docx, target, preview_path, ole_path, preview_width, preview_height, f"Equation.{mathtype_version}")
    return target


def _inject_object(
    input_docx: Path,
    output_docx: Path,
    preview_path: Path,
    ole_path: Path,
    preview_width_px: int,
    preview_height_px: int,
    prog_id: str = "Equation.DSMT4",
) -> None:
    parser = etree.XMLParser(resolve_entities=False, recover=True, remove_blank_text=False)
    with ZipFile(input_docx) as zin:
        document_root = etree.fromstring(zin.read("word/document.xml"), parser)
        rels_root = etree.fromstring(zin.read("word/_rels/document.xml.rels"), parser)
        content_types_root = etree.fromstring(zin.read("[Content_Types].xml"), parser)

        image_rel_id = _add_relationship(rels_root, IMAGE_REL, "media/mathtype_preview_001.png")
        ole_rel_id = _add_relationship(rels_root, OLE_REL, "embeddings/oleObject1.bin")
        _ensure_default(content_types_root, "png", "image/png")
        _ensure_default(content_types_root, "bin", "application/vnd.openxmlformats-officedocument.oleObject")

        placeholder_run = _find_placeholder_run(document_root)
        object_run = _object_run(image_rel_id, ole_rel_id, preview_width_px, preview_height_px, prog_id=prog_id)
        placeholder_run.getparent().replace(placeholder_run, object_run)

        document_xml = etree.tostring(document_root, encoding="utf-8", xml_declaration=True, standalone=True)
        rels_xml = etree.tostring(rels_root, encoding="utf-8", xml_declaration=True, standalone=True)
        content_types_xml = etree.tostring(content_types_root, encoding="utf-8", xml_declaration=True, standalone=True)

        with ZipFile(output_docx, "w", ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                if info.filename == "word/document.xml":
                    zout.writestr(info, document_xml)
                elif info.filename == "word/_rels/document.xml.rels":
                    zout.writestr(info, rels_xml)
                elif info.filename == "[Content_Types].xml":
                    zout.writestr(info, content_types_xml)
                else:
                    zout.writestr(info, zin.read(info.filename))
            zout.writestr("word/media/mathtype_preview_001.png", preview_path.read_bytes())
            zout.writestr("word/embeddings/oleObject1.bin", ole_path.read_bytes())


def _add_relationship(root: etree._Element, rel_type: str, target: str) -> str:
    existing = {rel.get("Id") for rel in root.findall(f"{{{REL_NS}}}Relationship")}
    index = 1
    while f"rIdMT{index}" in existing:
        index += 1
    rel_id = f"rIdMT{index}"
    etree.SubElement(
        root,
        f"{{{REL_NS}}}Relationship",
        {"Id": rel_id, "Type": rel_type, "Target": target},
    )
    return rel_id


def _ensure_default(root: etree._Element, extension: str, content_type: str) -> None:
    if not root.xpath(f'ct:Default[@Extension="{extension}"]', namespaces={"ct": CT_NS}):
        etree.SubElement(root, f"{{{CT_NS}}}Default", {"Extension": extension, "ContentType": content_type})


def _find_placeholder_run(root: etree._Element) -> etree._Element:
    for run in root.xpath("//w:r", namespaces=NS):
        text = "".join(run.xpath(".//w:t/text()", namespaces=NS))
        if PLACEHOLDER in text:
            return run
    raise ValueError("Placeholder run was not found in the generated DOCX.")


def make_object_run(
    image_rel_id: str,
    ole_rel_id: str,
    preview_width_px: int,
    preview_height_px: int,
    index: int = 1,
    display: bool = False,
    inline_height_pt: float = 12.5,
    display_height_pt: float = 21.0,
    max_width_pt: float = 360.0,
    preview_pt_per_px: float | None = None,
    vertical_align: str | None = None,
    prog_id: str = "Equation.DSMT4",
) -> etree._Element:
    width_pt, height_pt = _fit_preview_size(
        preview_width_px,
        preview_height_px,
        display=display,
        inline_height_pt=inline_height_pt,
        display_height_pt=display_height_pt,
        max_width_pt=max_width_pt,
        preview_pt_per_px=preview_pt_per_px,
    )
    width_dxa = int(width_pt * 20)
    height_dxa = int(height_pt * 20)
    shape_id = f"_x0000_i{1024 + index}"
    object_id = f"_{1618029000 + index}"

    run = etree.Element(_q("w:r"), nsmap=NS)
    obj = etree.SubElement(
        run,
        _q("w:object"),
        {
            _q("w:dxaOrig"): str(width_dxa),
            _q("w:dyaOrig"): str(height_dxa),
            _q("w14:anchorId"): f"{index:08X}"[-8:],
        },
    )
    shape_type = etree.SubElement(
        obj,
        _q("v:shapetype"),
        {
            "id": "_x0000_t75",
            "coordsize": "21600,21600",
            _q("o:spt"): "75",
            _q("o:preferrelative"): "t",
            "path": "m@4@5l@4@11@9@11@9@5xe",
            "filled": "f",
            "stroked": "f",
        },
    )
    etree.SubElement(shape_type, _q("v:stroke"), {"joinstyle": "miter"})
    formulas = etree.SubElement(shape_type, _q("v:formulas"))
    for equation in [
        "if lineDrawn pixelLineWidth 0",
        "sum @0 1 0",
        "sum 0 0 @1",
        "prod @2 1 2",
        "prod @3 21600 pixelWidth",
        "prod @3 21600 pixelHeight",
        "sum @0 0 1",
        "prod @6 1 2",
        "prod @7 21600 pixelWidth",
        "sum @8 21600 0",
        "prod @7 21600 pixelHeight",
        "sum @10 21600 0",
    ]:
        etree.SubElement(formulas, _q("v:f"), {"eqn": equation})
    etree.SubElement(shape_type, _q("v:path"), {_q("o:extrusionok"): "f", "gradientshapeok": "t", _q("o:connecttype"): "rect"})
    etree.SubElement(shape_type, _q("o:lock"), {_q("v:ext"): "edit", "aspectratio": "t"})

    shape_style = f"width:{width_pt:.1f}pt;height:{height_pt:.1f}pt"
    if vertical_align:
        shape_style += f";vertical-align:{vertical_align}"

    shape = etree.SubElement(
        obj,
        _q("v:shape"),
        {
            "id": shape_id,
            "type": "#_x0000_t75",
            "style": shape_style,
            _q("o:ole"): "",
        },
    )
    etree.SubElement(shape, _q("v:imagedata"), {_q("r:id"): image_rel_id, _q("o:title"): ""})
    etree.SubElement(
        obj,
        _q("o:OLEObject"),
        {
            "Type": "Embed",
            "ProgID": prog_id,
            "ShapeID": shape_id,
            "DrawAspect": "Content",
            "ObjectID": object_id,
            _q("r:id"): ole_rel_id,
        },
    )
    return run


def _object_run(
    image_rel_id: str,
    ole_rel_id: str,
    preview_width_px: int,
    preview_height_px: int,
    prog_id: str = "Equation.DSMT4",
) -> etree._Element:
    return make_object_run(image_rel_id, ole_rel_id, preview_width_px, preview_height_px, index=1, display=True, prog_id=prog_id)


def make_display_equation_paragraph(
    object_run: etree._Element,
    number: int | str,
    text_width_dxa: int,
    *,
    before_dxa: int = 80,
    after_dxa: int = 80,
    numbering: NumberingOptions | None = None,
    style: EquationStyle | None = None,
) -> etree._Element:
    return make_tabbed_equation_paragraph(
        object_run,
        number,
        text_width_dxa,
        numbering=numbering or NumberingOptions(before_dxa=before_dxa, after_dxa=after_dxa, use_seq_field=False),
        style=style or EquationStyle(),
    )


def _fit_preview_size(
    width_px: int,
    height_px: int,
    display: bool,
    inline_height_pt: float,
    display_height_pt: float,
    max_width_pt: float,
    preview_pt_per_px: float | None = None,
) -> tuple[float, float]:
    if height_px <= 0:
        return 20.0, inline_height_pt
    if preview_pt_per_px is not None:
        target_width = width_px * preview_pt_per_px
        target_height = height_px * preview_pt_per_px
    else:
        target_height = display_height_pt if display else inline_height_pt
        target_width = target_height * width_px / height_px
    if target_width > max_width_pt:
        scale = max_width_pt / target_width
        target_width = max_width_pt
        target_height *= scale
    return max(8.0, target_width), max(8.0, target_height)


def _tab_run() -> etree._Element:
    run = etree.Element(_q("w:r"), nsmap=NS)
    etree.SubElement(run, _q("w:tab"))
    return run


def _text_run(text: str) -> etree._Element:
    run = etree.Element(_q("w:r"), nsmap=NS)
    rpr = etree.SubElement(run, _q("w:rPr"))
    etree.SubElement(
        rpr,
        _q("w:rFonts"),
        {
            _q("w:ascii"): "Times New Roman",
            _q("w:hAnsi"): "Times New Roman",
            _q("w:eastAsia"): "SimSun",
        },
    )
    etree.SubElement(rpr, _q("w:sz"), {_q("w:val"): "21"})
    etree.SubElement(rpr, _q("w:szCs"), {_q("w:val"): "21"})
    text_el = etree.SubElement(run, _q("w:t"))
    text_el.text = text
    return run


def _q(name: str) -> str:
    prefix, local = name.split(":", 1)
    return f"{{{NS[prefix]}}}{local}"
