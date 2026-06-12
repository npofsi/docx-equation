from __future__ import annotations

from pathlib import Path
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree

from docx_equation.omml.converter import mathml_to_omml
from docx_equation.shared.models import ConversionSummary, EquationSpec, ExportOptions, OptionsLike, normalize_options
from docx_equation.shared.numbering import make_tabbed_equation_paragraph
from docx_equation.shared.ooxml import ancestor_paragraph, document_text_width_dxa, find_placeholder_run


def embed_mathml_placeholders(
    input_docx: str | Path,
    output_docx: str | Path,
    equations: list[EquationSpec],
    options: OptionsLike = None,
    work_dir: str | Path | None = None,
) -> ConversionSummary:
    opts = normalize_options(options, target="omml")
    source = Path(input_docx)
    target = Path(output_docx)
    target.parent.mkdir(parents=True, exist_ok=True)
    if work_dir is None:
        with tempfile.TemporaryDirectory(prefix="docx_equation_omml_") as tmp:
            return _embed(source, target, equations, opts, Path(tmp))
    return _embed(source, target, equations, opts, Path(work_dir))


def build_equation_docx(
    equations: list[str],
    output_docx: str | Path,
    options: OptionsLike = None,
) -> ConversionSummary:
    opts = normalize_options(options, target="omml")
    target = Path(output_docx)
    target.parent.mkdir(parents=True, exist_ok=True)
    specs: list[EquationSpec] = []
    with tempfile.TemporaryDirectory(prefix="docx_equation_omml_docx_") as tmp:
        base = Path(tmp) / "base.docx"
        doc = Document()
        for index, mathml in enumerate(equations, 1):
            placeholder = f"{{{{DOCX_EQ_{index:03d}}}}}"
            doc.add_paragraph().add_run(placeholder)
            specs.append(EquationSpec(placeholder=placeholder, mathml=mathml, display=True, number=index))
        doc.save(base)
        return embed_mathml_placeholders(base, target, specs, opts, Path(tmp) / "work")


def _embed(
    source: Path,
    target: Path,
    equations: list[EquationSpec],
    opts: ExportOptions,
    work_dir: Path,
) -> ConversionSummary:
    work_dir.mkdir(parents=True, exist_ok=True)
    summary = ConversionSummary(found=len(equations))
    parser = etree.XMLParser(resolve_entities=False, recover=True, remove_blank_text=False)
    with ZipFile(source) as zin:
        document_root = etree.fromstring(zin.read("word/document.xml"), parser)
        for index, equation in enumerate(equations, 1):
            try:
                placeholder_run = find_placeholder_run(document_root, equation.placeholder)
                display_without_number = equation.display and equation.number is None
                replacement = mathml_to_omml(equation.mathml, display=display_without_number, options=opts.omml)
                if equation.display and equation.number is not None and opts.display_layout == "tabbed":
                    paragraph = make_tabbed_equation_paragraph(
                        replacement,
                        equation.number,
                        document_text_width_dxa(document_root),
                        numbering=opts.numbering,
                        style=opts.style,
                    )
                    placeholder_paragraph = ancestor_paragraph(placeholder_run)
                    if placeholder_paragraph is None:
                        raise ValueError(f"Placeholder paragraph was not found: {equation.placeholder}")
                    paragraph.tail = placeholder_paragraph.tail
                    placeholder_paragraph.getparent().replace(placeholder_paragraph, paragraph)
                else:
                    replacement.tail = placeholder_run.tail
                    placeholder_run.getparent().replace(placeholder_run, replacement)
                summary.converted += 1
            except Exception as exc:  # noqa: BLE001 - conversion continues per equation.
                summary.add_error(index, equation.placeholder, str(exc))

        document_xml = etree.tostring(document_root, encoding="utf-8", xml_declaration=True, standalone=True)
        with ZipFile(target, "w", ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                if info.filename == "word/document.xml":
                    zout.writestr(info, document_xml)
                else:
                    zout.writestr(info, zin.read(info.filename))
    return summary
