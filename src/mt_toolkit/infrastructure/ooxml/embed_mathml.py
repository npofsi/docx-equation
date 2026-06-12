from __future__ import annotations

from pathlib import Path
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree
from PIL import Image, ImageDraw

from mt_toolkit.docx_embed import IMAGE_REL, NS, OLE_REL, _add_relationship, _ensure_default, make_object_run
from mt_toolkit.domain.models import ConversionOptions, ConversionSummary, EquationSpec
from mt_toolkit.infrastructure.ooxml.mathml_to_omml import mathml_to_omml
from mt_toolkit.mathml import render_mathml_files
from mt_toolkit.mtef import encode_mtef
from mt_toolkit.ole import build_mathtype_ole_object


REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
MTYPE_NS = "https://github.com/npofsi/mt-toolkit/mathtype"
W_NS = NS["w"]


def embed_mathml_placeholders(
    input_docx: str | Path,
    output_docx: str | Path,
    equations: list[EquationSpec],
    options: ConversionOptions | None = None,
    work_dir: str | Path | None = None,
) -> ConversionSummary:
    opts = options or ConversionOptions()
    source = Path(input_docx)
    target = Path(output_docx)
    target.parent.mkdir(parents=True, exist_ok=True)
    if work_dir is None:
        with tempfile.TemporaryDirectory(prefix="mt_toolkit_embed_") as tmp:
            return _embed(source, target, equations, opts, Path(tmp))
    return _embed(source, target, equations, opts, Path(work_dir))


def build_equation_docx(
    equations: list[str],
    output_docx: str | Path,
    options: ConversionOptions | None = None,
) -> ConversionSummary:
    opts = options or ConversionOptions()
    target = Path(output_docx)
    target.parent.mkdir(parents=True, exist_ok=True)
    specs: list[EquationSpec] = []
    with tempfile.TemporaryDirectory(prefix="mt_toolkit_docx_") as tmp:
        base = Path(tmp) / "base.docx"
        doc = Document()
        for index, mathml in enumerate(equations, 1):
            placeholder = f"{{{{MT_EQ_{index:03d}}}}}"
            paragraph = doc.add_paragraph()
            paragraph.add_run(placeholder)
            specs.append(EquationSpec(placeholder=placeholder, mathml=mathml, display=True))
        doc.save(base)
        return embed_mathml_placeholders(base, target, specs, opts, Path(tmp) / "work")


def _embed(
    source: Path,
    target: Path,
    equations: list[EquationSpec],
    opts: ConversionOptions,
    work_dir: Path,
) -> ConversionSummary:
    work_dir.mkdir(parents=True, exist_ok=True)
    mathml_dir = work_dir / "mathml"
    preview_dir = work_dir / "preview_png"
    ole_dir = work_dir / "ole"
    for directory in (mathml_dir, preview_dir, ole_dir):
        directory.mkdir(parents=True, exist_ok=True)

    summary = ConversionSummary(found=len(equations))
    parser = etree.XMLParser(resolve_entities=False, recover=True, remove_blank_text=False)
    with ZipFile(source) as zin:
        document_root = etree.fromstring(zin.read("word/document.xml"), parser)
        rels_root = etree.fromstring(zin.read("word/_rels/document.xml.rels"), parser)
        content_types_root = etree.fromstring(zin.read("[Content_Types].xml"), parser)

        _ensure_mc(document_root)
        _ensure_default(content_types_root, "bin", "application/vnd.openxmlformats-officedocument.oleObject")
        _ensure_default(content_types_root, "png", "image/png")

        for index, equation in enumerate(equations, 1):
            (mathml_dir / f"equation_{index:03d}.mml").write_text(equation.mathml, encoding="utf-8")

        _render_previews(mathml_dir, preview_dir, opts)

        media_entries: list[tuple[str, bytes]] = []
        ole_entries: list[tuple[str, bytes]] = []

        for index, equation in enumerate(equations, 1):
            try:
                placeholder_run = _find_placeholder_run(document_root, equation.placeholder)
                preview_path = preview_dir / f"equation_{index:03d}.png"
                width_px, height_px = Image.open(preview_path).size

                image_name = f"mathtype_preview_{index:03d}.png"
                ole_name = f"oleObjectMathType{index:03d}.bin"
                image_rel_id = _add_relationship(rels_root, IMAGE_REL, f"media/{image_name}")
                ole_rel_id = _add_relationship(rels_root, OLE_REL, f"embeddings/{ole_name}")

                mtef = encode_mtef(_parse_mathml_expr(equation.mathml), opts.mathtype_version)
                ole_bytes = build_mathtype_ole_object(mtef, opts.prog_id)
                ole_entries.append((f"word/embeddings/{ole_name}", ole_bytes))
                media_entries.append((f"word/media/{image_name}", preview_path.read_bytes()))
                (ole_dir / ole_name).write_bytes(ole_bytes)

                replacement = make_object_run(
                    image_rel_id,
                    ole_rel_id,
                    width_px,
                    height_px,
                    index=index,
                    display=equation.display,
                    inline_height_pt=opts.inline_height_pt,
                    display_height_pt=opts.display_height_pt,
                    max_width_pt=opts.max_width_pt,
                    preview_pt_per_px=opts.preview_pt_per_px,
                    vertical_align="middle" if equation.display else None,
                    prog_id=opts.prog_id,
                )
                if opts.embed_mode == "alternate-content":
                    fallback = mathml_to_omml(equation.mathml, display=equation.display)
                    replacement = _alternate_content(replacement, fallback)
                replacement.tail = placeholder_run.tail
                parent = placeholder_run.getparent()
                parent.replace(placeholder_run, replacement)
                summary.converted += 1
            except Exception as exc:  # noqa: BLE001 - conversion continues per equation.
                summary.add_error(index, equation.placeholder, str(exc))

        document_xml = etree.tostring(document_root, encoding="utf-8", xml_declaration=True, standalone=True)
        rels_xml = etree.tostring(rels_root, encoding="utf-8", xml_declaration=True, standalone=True)
        content_types_xml = etree.tostring(content_types_root, encoding="utf-8", xml_declaration=True, standalone=True)

        with ZipFile(target, "w", ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                if info.filename == "word/document.xml":
                    zout.writestr(info, document_xml)
                elif info.filename == "word/_rels/document.xml.rels":
                    zout.writestr(info, rels_xml)
                elif info.filename == "[Content_Types].xml":
                    zout.writestr(info, content_types_xml)
                else:
                    zout.writestr(info, zin.read(info.filename))
            for name, data in media_entries + ole_entries:
                zout.writestr(name, data)
    return summary


def _parse_mathml_expr(mathml: str):
    from mt_toolkit.mathml import parse_mathml

    return parse_mathml(mathml)


def _render_previews(mathml_dir: Path, preview_dir: Path, opts: ConversionOptions) -> None:
    try:
        render_mathml_files(mathml_dir, preview_dir, chrome_path=opts.chrome_path)
    except Exception:
        if opts.embed_mode == "png-preview":
            raise
        for mathml_path in sorted(mathml_dir.glob("equation_*.mml")):
            _write_fallback_preview(preview_dir / f"{mathml_path.stem}.png")


def _write_fallback_preview(path: Path) -> None:
    image = Image.new("RGB", (420, 96), "white")
    draw = ImageDraw.Draw(image)
    draw.text((16, 34), "MathType equation", fill="black")
    image.save(path)


def _find_placeholder_run(root: etree._Element, placeholder: str) -> etree._Element:
    for run in root.xpath("//w:r", namespaces=NS):
        text = "".join(run.xpath(".//w:t/text()", namespaces=NS))
        if text == placeholder:
            return run
    raise ValueError(f"Placeholder run was not found: {placeholder}")


def _alternate_content(choice_child: etree._Element, fallback_child: etree._Element) -> etree._Element:
    alternate = etree.Element(f"{{{MC_NS}}}AlternateContent", nsmap={"mc": MC_NS, "mtype": MTYPE_NS})
    choice = etree.SubElement(alternate, f"{{{MC_NS}}}Choice", {"Requires": "mtype"})
    choice.append(choice_child)
    fallback = etree.SubElement(alternate, f"{{{MC_NS}}}Fallback")
    fallback.append(fallback_child)
    return alternate


def _ensure_mc(root: etree._Element) -> None:
    ignorable_attr = f"{{{MC_NS}}}Ignorable"
    current = root.get(ignorable_attr, "")
    values = {item for item in current.split() if item}
    values.add("mtype")
    root.set(ignorable_attr, " ".join(sorted(values)))
