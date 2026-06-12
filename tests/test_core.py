from __future__ import annotations

from zipfile import ZipFile

from lxml import etree
import olefile
from docx import Document

from docx_equation import (
    ConversionOptions,
    EquationStyle,
    EquationSpec,
    ExportOptions,
    NumberingOptions,
    embed_mathml_placeholders,
    inspect_docx,
    mathml_to_mathtype_ole,
    mathml_to_omml,
    make_tabbed_equation_paragraph,
)
from docx_equation.mathtype.ooxml import NS, build_demo_docx, make_display_equation_paragraph
from docx_equation.shared.latex import parse_latex_subset
from docx_equation.shared.mathml import parse_mathml
from docx_equation.mathtype.mtef import encode_mtef
from docx_equation.mathtype.ole import build_mathtype_ole_object


def test_ole_contains_mathtype_streams(tmp_path):
    expr = parse_latex_subset(r"\frac{4}{5}=0.8")
    ole_path = tmp_path / "oleObject1.bin"
    ole_path.write_bytes(build_mathtype_ole_object(encode_mtef(expr)))

    ole = olefile.OleFileIO(str(ole_path))
    try:
        streams = {"/".join(item) for item in ole.listdir(streams=True, storages=False)}
        assert "\x01CompObj" in streams
        assert "\x01Ole" in streams
        assert "\x03ObjInfo" in streams
        assert "Equation Native" in streams
        native = ole.openstream(["Equation Native"]).read()
        assert native[28] == 5
    finally:
        ole.close()


def test_demo_docx_contains_equation_dsmt4_object(tmp_path):
    docx_path = tmp_path / "demo.docx"
    build_demo_docx(r"F_n=F_{n-1}+F_{n-2}", docx_path)

    with ZipFile(docx_path) as zf:
        names = set(zf.namelist())
        document_xml = zf.read("word/document.xml").decode("utf-8")
        rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "word/embeddings/oleObject1.bin" in names
        assert "word/media/mathtype_preview_001.png" in names
        assert 'ProgID="Equation.DSMT4"' in document_xml
        assert "relationships/oleObject" in rels_xml


def test_mathml_parser_supports_section_formula_subset():
    expr = parse_mathml(
        b"""<math xmlns="http://www.w3.org/1998/Math/MathML">
        <msubsup><mi>P</mi><mrow><mi>e</mi><mi>s</mi></mrow><mrow><mi>c</mi><mi>h</mi></mrow></msubsup>
        <mo>(</mo><mi>t</mi><mo>)</mo><mo>=</mo>
        <mfrac><mrow><mn>1</mn></mrow><mrow><mn>1</mn><mo>+</mo><mi>\xcf\x84</mi><mi>s</mi></mrow></mfrac>
        </math>"""
    )
    mtef = encode_mtef(expr)
    assert mtef.startswith(b"\x05\x01\x00\x06")
    assert b"DSMT4" in mtef
    assert b"DSMT6" in encode_mtef(expr, "DSMT6")


def test_latex_parser_supports_big_operator_and_decoration():
    expr = parse_latex_subset(r"\sum_{i=1}^{n} \overline{x_i}+\int_0^T p(t)")
    mtef = encode_mtef(expr)
    assert mtef.startswith(b"\x05\x01\x00\x06")
    assert b"\x10" in mtef
    assert b"\x0f" in mtef


def test_mathml_parser_supports_matrix_and_limits():
    expr = parse_mathml(
        b"""<math xmlns="http://www.w3.org/1998/Math/MathML">
        <munderover><mo>&#x2211;</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></munderover>
        <mtable><mtr><mtd><mi>a</mi></mtd><mtd><mi>b</mi></mtd></mtr><mtr><mtd><mi>c</mi></mtd><mtd><mi>d</mi></mtd></mtr></mtable>
        </math>"""
    )
    mtef = encode_mtef(expr)
    assert mtef.startswith(b"\x05\x01\x00\x06")
    assert b"\x05\x00" in mtef


def test_display_equation_paragraph_uses_center_and_right_tabs():
    object_run = etree.Element(f"{{{NS['w']}}}r", nsmap=NS)
    paragraph = make_display_equation_paragraph(object_run, "(4-1)", 8958)
    tabs = paragraph.xpath(".//w:tab[@w:val]", namespaces=NS)
    assert [tab.get(f"{{{NS['w']}}}val") for tab in tabs] == ["center", "right"]
    assert [tab.get(f"{{{NS['w']}}}pos") for tab in tabs] == ["4479", "8958"]
    assert "(4-1)" in "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))


def test_numbering_options_support_sequence_only_format():
    object_run = etree.Element(f"{{{NS['w']}}}r", nsmap=NS)
    paragraph = make_tabbed_equation_paragraph(
        object_run,
        1,
        8958,
        numbering=NumberingOptions(chapter=4, number_format="(1)"),
        style=EquationStyle(),
    )
    document_xml = etree.tostring(paragraph, encoding="unicode")
    text = "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))

    assert text == "(1)"
    assert "(4-" not in text
    assert " SEQ Eq \\r 1 \\* ARABIC " in document_xml


def test_numbering_options_support_chapter_sequence_separator():
    object_run = etree.Element(f"{{{NS['w']}}}r", nsmap=NS)
    paragraph = make_tabbed_equation_paragraph(
        object_run,
        2,
        8958,
        numbering=NumberingOptions(chapter=4, number_format="(1SEP1)", separator="."),
        style=EquationStyle(),
    )
    document_xml = etree.tostring(paragraph, encoding="unicode")
    text = "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))

    assert text == "(4.2)"
    assert " SEQ Eq \\* ARABIC " in document_xml


def test_mathml_to_mathtype_ole_uses_dsmt4_by_default(tmp_path):
    mathml = """<math xmlns="http://www.w3.org/1998/Math/MathML">
    <mfrac><mi>a</mi><mi>b</mi></mfrac>
    </math>"""
    ole_path = tmp_path / "equation.bin"
    ole_path.write_bytes(mathml_to_mathtype_ole(mathml))

    ole = olefile.OleFileIO(str(ole_path))
    try:
        native = ole.openstream(["Equation Native"]).read()
        comp = ole.openstream(["\x01CompObj"]).read()
        assert b"DSMT4" in native
        assert b"Equation.DSMT4" in comp
    finally:
        ole.close()


def test_mathml_to_omml_returns_word_math():
    mathml = """<math xmlns="http://www.w3.org/1998/Math/MathML">
    <msub><mi>P</mi><mtext>es</mtext></msub><mo>=</mo><mn>1</mn>
    </math>"""
    omml = mathml_to_omml(mathml)
    assert etree.QName(omml).localname == "oMath"
    assert omml.xpath(".//m:sSub", namespaces={"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"})


def test_embed_mathml_placeholders_adds_alternate_content_and_ole(tmp_path):
    base = tmp_path / "base.docx"
    output = tmp_path / "output.docx"
    placeholder = "{{DOCX_EQ_001}}"
    mathml = """<math xmlns="http://www.w3.org/1998/Math/MathML">
    <msup><mi>x</mi><mn>2</mn></msup><mo>+</mo><mn>1</mn>
    </math>"""

    doc = Document()
    doc.add_paragraph().add_run(placeholder)
    doc.save(base)

    summary = embed_mathml_placeholders(
        base,
        output,
        [EquationSpec(placeholder=placeholder, mathml=mathml, display=False)],
        ConversionOptions(embed_mode="alternate-content"),
        tmp_path / "work",
    )
    assert summary.converted == 1

    with ZipFile(output) as zf:
        names = set(zf.namelist())
        document_root = etree.fromstring(zf.read("word/document.xml"))
        document_xml = etree.tostring(document_root, encoding="unicode")
        assert any(name.startswith("word/embeddings/oleObjectMathType") for name in names)
        assert "AlternateContent" in document_xml
        assert "Equation.DSMT4" in document_xml
        assert document_root.nsmap["dxeq"] == "https://github.com/npofsi/docx-equation/mathtype"
        assert "dxeq" in document_root.get("{http://schemas.openxmlformats.org/markup-compatibility/2006}Ignorable").split()
        assert placeholder not in document_xml
    counts = inspect_docx(output)
    assert counts["alternate_content"] == 1
    assert counts["embeddings"] == 1


def test_embed_mathml_placeholders_can_emit_omml_with_tabbed_numbering(tmp_path):
    base = tmp_path / "base.docx"
    output = tmp_path / "output.docx"
    placeholder = "{{DOCX_EQ_001}}"
    mathml = """<math xmlns="http://www.w3.org/1998/Math/MathML" display="block">
    <mfrac><mi>a</mi><mi>b</mi></mfrac><mo>=</mo><mi>c</mi>
    </math>"""

    doc = Document()
    doc.add_paragraph().add_run(placeholder)
    doc.save(base)

    summary = embed_mathml_placeholders(
        base,
        output,
        [EquationSpec(placeholder=placeholder, mathml=mathml, display=True, number=1)],
        ExportOptions(target="omml", numbering=NumberingOptions(chapter=4)),
        tmp_path / "work",
    )
    assert summary.converted == 1

    with ZipFile(output) as zf:
        names = set(zf.namelist())
        document_root = etree.fromstring(zf.read("word/document.xml"))
        document_xml = etree.tostring(document_root, encoding="unicode")
        tabs = document_root.xpath(".//w:tab[@w:val]", namespaces=NS)
        assert not any(name.startswith("word/embeddings/") for name in names)
        assert document_root.xpath(".//m:oMath", namespaces={"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"})
        assert [tab.get(f"{{{NS['w']}}}val") for tab in tabs[:2]] == ["center", "right"]
        assert " SEQ Eq \\r 1 \\* ARABIC " in document_xml
        assert placeholder not in document_xml
